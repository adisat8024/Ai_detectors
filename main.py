import streamlit as st
import pandas as pd
import os
import numpy as np
import tensorflow as tf
import cv2
import PyPDF2
import docx
import pdfplumber
import pickle
import joblib
import requests
from bs4 import BeautifulSoup
from docx import Document
from tensorflow.keras.models import load_model
from spacy.lang.en.stop_words import STOP_WORDS as stopwords
import string

FEEDBACK_FILE = "feedback.csv"
REVIEWS_FILE = "user_reviews.csv"

# Function to save user review
def save_review(name, review, rating):
    review_entry = pd.DataFrame([[name, review, rating]], columns=["Name", "Review", "Rating"])
    review_entry.to_csv(REVIEWS_FILE, mode='a', header=not os.path.exists(REVIEWS_FILE), index=False)

# Load models once at startup
@st.cache_resource
def load_ai_image_model():
    return load_model("ai_detector_model.keras")

@st.cache_resource
def load_text_ai_model():
    with open('clf.pkl', 'rb') as clf_file, open('tfidf.pkl', 'rb') as tfidf_file:
        return pickle.load(clf_file), pickle.load(tfidf_file)

@st.cache_resource
def load_fake_news_model():
    return joblib.load("vectorizer.jb"), joblib.load("lr.jb")

# Load models
image_model = load_ai_image_model()
text_clf, text_tfidf = load_text_ai_model()
news_vectorizer, news_model = load_fake_news_model()

def read_file_content(uploaded_file):
    """Reads the content of a text, PDF, DOC, or DOCX file."""
    file_extension = uploaded_file.name.split(".")[-1].lower()

    if file_extension == "txt":
        return uploaded_file.read().decode("utf-8")
    
    elif file_extension == "pdf":
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        text = "\n".join(page.extract_text() for page in pdf_reader.pages if page.extract_text())
        return text
    
    elif file_extension in ["doc", "docx"]:
        doc = docx.Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    
    else:
        st.error("Unsupported file format! Please upload a .txt, .pdf, .doc, or .docx file.")
        return None

if "theme" not in st.session_state:
    st.session_state["theme"] = "Dark"

# Sidebar: Theme Selection
theme_selection = st.sidebar.radio("Choose Theme", ["Dark", "Light"])
# Sidebar Navigation
st.sidebar.title("🔍 Choose a Detector")
app_mode = st.sidebar.radio("Select a Detection Mode", ["AI Image Detector", "AI Text Detector", "Misinformation Detector"])

if "last_app_mode" not in st.session_state:
    st.session_state["last_app_mode"] = app_mode

if st.session_state["last_app_mode"] != app_mode:
    # Clear relevant session state variables when mode changes
    if "text_result" in st.session_state:
        del st.session_state["text_result"]
    if "feedback_text" in st.session_state:
        del st.session_state["feedback_text"]
    if "news_result" in st.session_state:
        del st.session_state["news_result"]
    if "feedback_news" in st.session_state:
        del st.session_state["feedback_news"]
    if "input_text" in st.session_state:
        del st.session_state["input_text"]
    if "last_method" in st.session_state:
        del st.session_state["last_method"]
    if "feedback" in st.session_state:
        del st.session_state["feedback"]
    st.session_state["last_app_mode"] = app_mode

if theme_selection == "Dark":
    primary_color = "#F63366"
    background_color = "#0E1117"
    secondary_background_color = "#262730"
    text_color = "#FFFFFF"
else:
    primary_color = "#3366FF"
    background_color = "#FFFFFF"
    secondary_background_color = "#FFFFFF"
    text_color = "#000000"

# Apply custom styling with higher specificity
st.markdown(
    f"""
    <style>
        /* Main app and body */
        body, .stApp {{
            background-color: {background_color} !important;
            color: {text_color} !important;
        }}

        /* Sidebar */
        .stSidebar, .stSidebar .sidebar-content {{
            background-color: {secondary_background_color} !important;
            color: {text_color} !important;
        }}
        .stSidebar label, .stSidebar span {{
            color: {text_color} !important;
        }}

        /* File Uploader */
        [data-testid="stFileUploaderDropzone"] {{
            background-color: {background_color} !important;
            color: {text_color} !important;
        }}
        [data-testid="stFileUploaderDropzoneInstructions"] span, 
        [data-testid="stFileUploaderDropzoneInstructions"] small {{
            color: {text_color} !important;
        }}
        [data-testid="stBaseButton-secondary"] {{
            background-color: {background_color} !important;
            color: {text_color} !important;
            border: 1px solid {text_color} !important;
        }}

        /* Text Area/Input (assumed structure, refine with inspection) */
        [data-testid="stTextArea"] textarea, [data-testid="stTextInput"] textarea {{
            background-color: {background_color} !important;
            color: {text_color} !important;
            border: 1px solid {text_color} !important;
        }}
        [data-testid="stTextArea"] label, [data-testid="stTextInput"] label {{
            color: {text_color} !important;
        }}
        /* Text Input (st.text_input) */
        [data-testid="stTextInput"] input {{
            background-color: {background_color} !important;
            color: {text_color} !important;
            border: 1px solid {text_color} !important;
        }}
        [data-testid="stTextInput"] label {{
            color: {text_color} !important;
        }}
        /* Header */
        [data-testid="stHeader"] {{
            background-color: {secondary_background_color} !important;
            color: {text_color} !important;
        }}
        [data-testid="stDecoration"] {{
            background-color: {secondary_background_color} !important;
        }}
        [data-testid="stToolbar"] {{
            background-color: {secondary_background_color} !important;
            color: {text_color} !important;
        }}
        [data-testid="stBaseButton-header"], [data-testid="stBaseButton-headerNoPadding"] {{
            background-color: {secondary_background_color} !important;
            color: {text_color} !important;
            border: 1px solid {text_color} !important;
        }}
        [data-testid="stBaseButton-header"] span, [data-testid="stBaseButton-headerNoPadding"] span {{
            color: {text_color} !important;
        }}
        [data-testid="stBaseButton-header"] svg, [data-testid="stBaseButton-headerNoPadding"] svg {{
            fill: {text_color} !important;
        }}
        /* Dropdown Options (Menu Items) */
        ul[role="option"], ul[role="option"] li, ul[role="option"] span {{
            background-color: {secondary_background_color} !important;
            color: {text_color} !important;
        }}
        ul[role="option"] li:hover {{
            background-color: {primary_color} !important; /* Optional: Highlight on hover */
            color: {text_color} !important;
        }}
        [data-testid="stMainMenuDivider"] {{
            background-color: {secondary_background_color} !important;
        }}
        div[class*="st-emotion-cache-"] {{
            color: {text_color} !important;
        }}
        /* Buttons */
        .stButton > button {{
            background-color: {primary_color} !important;
            color: white !important;
        }}

        /* General text elements */
        h1, h2, h3, h4, h5, h6, p, label, span {{
            color: {text_color} !important;
        }}
    </style>
    """,
    unsafe_allow_html=True,
)

# Store theme selection in session state
st.session_state["theme"] = theme_selection


# 📌 **AI Image Detector**
if app_mode == "AI Image Detector":
    st.title("🖼️ AI-Generated Image Detector")
    uploaded_files = st.file_uploader("Upload images...", type=["jpg", "png", "jpeg"], accept_multiple_files=True)

    if uploaded_files:
        for uploaded_file in uploaded_files:
            file_bytes = np.asarray(bytearray(uploaded_file.read()), dtype=np.uint8)
            image = cv2.imdecode(file_bytes, 1)
            st.image(image, caption="Uploaded Image", use_column_width=True)

            with st.spinner("Analyzing image... Please wait."):
            # Preprocess Image
                image_resized = cv2.resize(image, (32, 32)) / 255.0
                processed_image = np.expand_dims(image_resized, axis=0)

            # Prediction
            prediction = image_model.predict(processed_image)
            ai_percentage = round(float(prediction[0][0]) * 100, 2)
            real_percentage = 100 - ai_percentage

            if ai_percentage > 50:
                result_text = f"🚨 Your image is AI-generated."
                st.error(result_text)
            else:
                result_text = f"✅ Your image is Real."
                st.success(result_text)

            # User Feedback
            if "feedback" not in st.session_state:
                st.session_state["feedback"] = {}
                
            # User selects feedbacr
            feedback = st.radio(f"Was this detection correct?", ["Yes", "No"], key=f"feedback_{uploaded_file.name}")
            
            st.session_state["feedback"][uploaded_file.name] = feedback
            if st.button(f"Submit Feedback"):
                selected_feedback = st.session_state["feedback"].get(uploaded_file.name)
                
                if selected_feedback:  # Ensures user selects an option
                    feedback_df = pd.DataFrame([['AI Image Detection', result_text, selected_feedback]], 
                                columns=["File", "Result", "Feedback"])
                    feedback_df.to_csv(FEEDBACK_FILE, mode='a', header=not os.path.exists(FEEDBACK_FILE), index=False)
                    st.success("Feedback Submitted! ✅")
            
            if "detection_result" in st.session_state:  # Ensure a detector was used
                st.session_state["detection_result"] = {}
    
            st.write("### 📝 Leave a Review")
            # Collect user input
            user_name = st.text_input("Enter your Name")
            ai_rating = st.slider("Rate the AI's Performance (1-5 Stars)", 1, 5, 5)
            website_rating = st.slider("Rate the Website Experience (1-5 Stars)", 1, 5, 5)
            review_text = st.text_area("Share your thoughts on the AI's prediction")
            website_feedback = st.text_area("How can we improve the website?")

            # Submit Review Button
            if st.button("Submit Review"):
                if user_name.strip() and review_text.strip() and website_feedback.strip():
                # Prepare review data
                    review_entry = {
                    "Name": user_name.strip(),
                    "AI Rating": ai_rating,
                    "Website Rating": website_rating,
                    "Review": review_text.strip(),
                    "Website Feedback": website_feedback.strip()
                     }
            
                    # Save to CSV
                    review_df = pd.DataFrame([review_entry], columns=["Name", "AI Rating", "Website Rating", "Review", "Website Feedback"])
                    review_df.to_csv("user_reviews.csv", mode='a', header=not os.path.exists("user_reviews.csv"), index=False)

                    st.success("✅ Review Submitted Successfully!")
                else:
                    st.warning("⚠️ Please fill in all fields before submitting.")
                
# 📌 **AI Text Detector**
elif app_mode == "AI Text Detector":
    st.title("📝 AI Text Detector")
    uploaded_file = st.file_uploader("Upload a document", type=["txt", "pdf", "doc", "docx"])

    # Show text area only if no file is uploaded
    if uploaded_file:
        with st.spinner('Extracting Information'):
            extracted_text = read_file_content(uploaded_file)
            if extracted_text:
                st.text_area("Extracted Text:", value=extracted_text, height=200, disabled=True)
    else:
        input_text = st.text_area("Enter text manually:", height=200)
    
    if st.button("Check AI Text"):
        text_to_analyze = extracted_text if uploaded_file else input_text
        if text_to_analyze:
            with st.spinner("Analyzing text... Please wait."):
                contractions = {"ain't": "am not", "aren't": "are not", "can't": "cannot", "can't've": "cannot have",
                            "could've": "could have", "couldn't": "could not", "didn't": "did not",
                            "doesn't": "does not", "don't": "do not", "hadn't": "had not", "hasn't": "has not",
                            "haven't": "have not", "he'd": "he would", "he'll": "he will", "he's": "he is",
                            "i'd": "i would", "i'll": "i will", "i'm": "i am", "i've": "i have", "isn't": "is not",
                            "it's": "it is", "let's": "let us", "mustn't": "must not", "shan't": "shall not",
                            "she'd": "she would", "she'll": "she will", "she's": "she is", "should've": "should have",
                            "shouldn't": "should not", "that's": "that is", "there's": "there is",
                            "they'd": "they would", "they'll": "they will", "they're": "they are", "they've": "they have",
                            "wasn't": "was not", "won't": "would not", "would've": "would have", "wouldn't": "would not",
                            " u ": " you ", " ur ": " your ", " n ": " and ", "dis": "this", "bak": "back", "brng": "bring"
                            }
            # Preprocess Text
                def clean_text(text):
                    tags = ['\n', '\'']
                    for tag in tags:
                        text = text.replace(tag, '')
                    text = [x for x in text if x not in string.punctuation]
                    text = ''.join(text)
                    text = ' '.join([t for t in text.split() if t.lower() not in stopwords])
                
                    if isinstance(text, str):
                        for key, value in contractions.items():
                            text = text.replace(key, value)
                    return text

                cleaned_text = clean_text(text_to_analyze)
                transformed_text = text_tfidf.transform([cleaned_text])
            result = text_clf.predict(transformed_text)
            result_text = "🚨 AI-Generated" if result[0] else "✅ Human-Written"

            if result[0]:
                st.error("🚨 Your text is AI-generated.")
            else:
                st.success("✅ Your text is human-written.")
            # Store result in session state to persist it
            st.session_state["text_result"] = result_text
            st.spinner("Analyis Completed")
        else:
            st.warning("Please enter some text before running the detector.")

    # Feedback section (only show if result exists)
    if "text_result" in st.session_state:
        if "feedback_text" not in st.session_state:
            st.session_state["feedback_text"] = None
        
        # Use a stable key unrelated to input_text
        feedback = st.radio("Was this detection correct?", ["Yes", "No"], key="text_feedback_radio")
        st.session_state["feedback_text"] = feedback

        if st.button("Submit Feedback"):
            with st.spinner("Submitting Feedback"):
                selected_feedback = st.session_state["feedback_text"]
                if selected_feedback:
                    feedback_df = pd.DataFrame([["AI Text Detection", st.session_state["text_result"], selected_feedback]], 
                                           columns=["Type", "Result", "Feedback"])
                    feedback_df.to_csv(FEEDBACK_FILE, mode='a', header=not os.path.exists(FEEDBACK_FILE), index=False)
                    st.success("Feedback Submitted! ✅")
                    # Clear feedback after submission
                    del st.session_state["text_result"]
                    del st.session_state["feedback_text"]


        st.write("### 📝 Leave a Review")
    
        # User Inputs
        user_name = st.text_input("Enter your Name")
        ai_rating = st.slider("Rate the AI's Performance (1-5 Stars)", 1, 5, 5)
        website_rating = st.slider("Rate the Website Experience (1-5 Stars)", 1, 5, 5)
        review_text = st.text_area("Share your thoughts on the AI's prediction")
        website_feedback = st.text_area("How can we improve the website?")

        # Submit Review Button
        if st.button("Submit Review"):
            if user_name.strip() and review_text.strip() and website_feedback.strip():
                # Save Review
                review_entry = {
                    "Name": user_name.strip(),
                    "AI Rating": ai_rating,
                    "Website Rating": website_rating,
                    "Review": review_text.strip(),
                    "Website Feedback": website_feedback.strip()
                }
                review_df = pd.DataFrame([review_entry])
                review_df.to_csv(REVIEWS_FILE, mode='a', header=not os.path.exists(REVIEWS_FILE), index=False)

                st.success("✅ Review Submitted Successfully!")
                del st.session_state["news_result"]  # Reset after submission
            else:
                st.warning("⚠️ Please fill in all fields before submitting.")

    


# 📌 **Misinformation Detector**
elif app_mode == "Misinformation Detector":
    st.title("📰 Fake News Detector")
    input_method = st.radio("Choose input method:", ["Enter Text", "Upload File", "Enter News URL"])
    
    if "input_text" not in st.session_state or st.session_state.get("last_method") != input_method:
        st.session_state["input_text"] = ""
        st.session_state["last_method"] = input_method
    # 🔹 Option 1: Manual Text Entry
    if input_method == "Enter Text":
        input_text = st.text_area("Enter news text manually:")
        st.session_state["input_text"] = input_text

    # 🔹 Option 2: File Upload (.txt, .pdf, .docx, .doc)
    elif input_method == "Upload File":
        uploaded_file = st.file_uploader("Upload a file (.txt, .pdf, .doc, .docx)", type=["txt", "pdf", "docx", "doc"])

        if uploaded_file:
            with st.spinner("Fetching Information"):
                file_extension = uploaded_file.name.split(".")[-1].lower()

                if file_extension == "pdf":
                    with pdfplumber.open(uploaded_file) as pdf:
                        extracted_text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())

                elif file_extension in ["docx", "doc"]:
                    doc = Document(uploaded_file)
                    extracted_text = "\n".join([para.text for para in doc.paragraphs])

                elif file_extension == "txt":
                    extracted_text = uploaded_file.read().decode("utf-8")
                st.session_state["input_text"] = extracted_text


# 🔹 Option 3: Extract News from URL
    elif input_method == "Enter News URL":
        news_url = st.text_input("Enter the news article URL:")
    
        if st.button("Fetch Article"):
            if news_url:
                with st.spinner("Fetching Article"):
                    try:
                        response = requests.get(news_url)
                        soup = BeautifulSoup(response.text, "html.parser")

                        # Extract the main article text
                        paragraphs = soup.find_all("p")
                        extracted_text = "\n".join([para.get_text() for para in paragraphs if para.get_text()])
                
                        if extracted_text.strip():
                            st.success("📰 Article text extracted successfully!")
                            st.session_state["input_text"] = extracted_text  # Store extracted text in session state
                        else:
                            st.error("⚠️ Failed to extract text. Try another URL.")
                            st.session_state["input_text"] = None  # Explicitly set to None

                    except Exception as e:
                        st.error(f"⚠️ Error fetching the article: {e}")
                        st.session_state["input_text"] = None 
    
    if input_method in ["Upload File", "Enter News URL"] and st.session_state["input_text"]:
        st.text_area("Extracted Article:", value=st.session_state["input_text"], height=200, disabled=True)

    if st.button("Check News"):
        input_text = st.session_state.get("input_text", "").strip()
        if input_text:
            with st.spinner("Checking news authenticity... Please wait."):
                transformed_input = news_vectorizer.transform([input_text])
                prediction = news_model.predict(transformed_input)
            result_text = "✅ The news is Real!" if prediction[0] == 1 else "🚨 The news is Fake!"

            if prediction[0] == 1:
                st.success("✅ The news is Real!")
            else:
                st.error("🚨 The news is Fake!")
            
            # Store result in session state
            st.session_state["news_result"] = result_text
        else:
            st.warning("⚠️ Please enter some text to analyze.")

    # Feedback section (only show if result exists)
    if "news_result" in st.session_state:
        if "feedback_news" not in st.session_state:
            st.session_state["feedback_news"] = None
        
        # Use a stable key unrelated to inputn
        feedback = st.radio("Was this detection correct?", ["Yes", "No"], key="news_feedback_radio")
        st.session_state["feedback_news"] = feedback

        if st.button("Submit Feedback"):
            with st.spinner("Submitting Feedback"):
                selected_feedback = st.session_state["feedback_news"]
                if selected_feedback:
                    feedback_df = pd.DataFrame([["Fake News Detection", st.session_state["news_result"], selected_feedback]], 
                                           columns=["Type", "Result", "Feedback"])
                    feedback_df.to_csv(FEEDBACK_FILE, mode='a', header=not os.path.exists(FEEDBACK_FILE), index=False)
                    st.success("Feedback Submitted! ✅")

                # Clear feedback after submission
                    del st.session_state["news_result"]
                    del st.session_state["feedback_news"]
                    
        st.write("### 📝 Leave a Review")
    
        # User Inputs
        user_name = st.text_input("Enter your Name")
        ai_rating = st.slider("Rate the AI's Performance (1-5 Stars)", 1, 5, 5)
        website_rating = st.slider("Rate the Website Experience (1-5 Stars)", 1, 5, 5)
        review_text = st.text_area("Share your thoughts on the AI's prediction")
        website_feedback = st.text_area("How can we improve the website?")

        # Submit Review Button
        if st.button("Submit Review"):
            if user_name.strip() and review_text.strip() and website_feedback.strip():
                # Save Review
                review_entry = {
                    "Name": user_name.strip(),
                    "AI Rating": ai_rating,
                    "Website Rating": website_rating,
                    "Review": review_text.strip(),
                    "Website Feedback": website_feedback.strip()
                }
                review_df = pd.DataFrame([review_entry])
                review_df.to_csv(REVIEWS_FILE, mode='a', header=not os.path.exists(REVIEWS_FILE), index=False)

                st.success("✅ Review Submitted Successfully!")
                del st.session_state["news_result"]  # Reset after submission
            else:
                st.warning("⚠️ Please fill in all fields before submitting.")


# Sidebar Button to View Feedback Data
if st.sidebar.button("View Feedback Data"):
    if os.path.exists(FEEDBACK_FILE) and os.path.getsize(FEEDBACK_FILE) > 0:
        feedback_data = pd.read_csv(FEEDBACK_FILE)
        with st.spinner("Getting Feedback"):

            if not feedback_data.empty:
                st.write("### 📊 User Feedback Data")
                st.dataframe(feedback_data)

                # 🔹 Calculate Agreement Percentage
                total_feedback = len(feedback_data)
                agreed_count = feedback_data[feedback_data["Feedback"] == "Yes"].shape[0]
                disagreement_count = feedback_data[feedback_data["Feedback"] == "No"].shape[0]

                agreement_percentage = (agreed_count / total_feedback) * 100 if total_feedback > 0 else 0
                disagreement_percentage = (disagreement_count / total_feedback) * 100 if total_feedback > 0 else 0

                # 🔹 Display Feedback Statistics
                st.write("### 📈 Feedback Summary")
                st.metric(label="Users Agreed with AI's Prediction", value=f"{agreement_percentage:.2f}%")
                st.metric(label="Users Disagreed with AI's Prediction", value=f"{disagreement_percentage:.2f}%")

            else:
                st.warning("⚠️ No feedback data found. Please submit feedback first!")

    else:
        st.warning("⚠️ No feedback data found. Please submit feedback first!")

                
# View User Reviews
if st.sidebar.button("View User Reviews"):
    if os.path.exists("user_reviews.csv") and os.path.getsize("user_reviews.csv") > 0:
        # Load the reviews file
        reviews_data = pd.read_csv("user_reviews.csv")

        # Check if data exists
        if not reviews_data.empty:
            st.write("### 📢 User Reviews")

            # Loop through each review and display
            for _, row in reviews_data.iterrows():
                st.write(f"👤 **Name:** {row['Name']}")
                st.write(f"🌟 **AI Rating:** {row['AI Rating']} / 5")
                st.write(f"💻 **Website Rating:** {row['Website Rating']} / 5")
                st.write(f"📝 **Review:** {row['Review']}")
                st.write(f"💡 **Suggestions for Improvement:** {row['Website Feedback']}")
                st.markdown("---")  # Separator between reviews
        else:
            st.warning("⚠️ No user reviews available.")
    else:
        st.warning("⚠️ No user reviews found. Please submit a review first!")
